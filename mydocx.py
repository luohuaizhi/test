# -*- coding:utf-8 -*-
import docx
import re
import traceback
import time


def cnt_time(fn):
    def wrap(*args):
        stime = time.time()
        fn(*args)
        runtime = time.time() - stime
        print runtime
    return wrap


order_id="11111111111111"
name = "test"
id_card_number = "2345678234567"
lender = "张三"
lender_id = "520555164502360052"
borrower = "成都银行"
borrower_id = "520555164502396562"
amount = "10000.00"
repay_type = "借记卡"
repay_times = "12"
repay_start_year = "2017"
repay_start_month = "08"
repay_start_day = "16"
repay_end_year = "2018"
repay_end_month = "07"
repay_end_day = "16"
repay_day_of_month = "16"
service_fee = "100"
insurance = "10000.000"
protect_plan = "10000.00"
bank_username = "张三"
bank_number = "8520741096374141052"
bank_branch_name = "成都银行高新区分行"
year, month, day = "2017", "08", "18"
sign_date = "2017年08月18日"
safe_fee = "12"
repay_count = "24"
user_address = "四川成都"
phone = "12345678911"
interest_rate = "信用卡扣"
start_year = "2017"
start_month = "08"
start_day = "09"
end_year = "2018"
end_month = "07"
end_day = "10"
bank_card_name = name
bank_card_id = "85274967485"


@cnt_time
def main(file_path):
    
    document = docx.Document(file_path)
    # 读取每段资料
    for paragraph in document.paragraphs:
        content = paragraph.text
        properties = get_param(content)
        if properties:
            for prop in properties:
                try:
                    content = content.replace("{"+prop+"}", eval(prop))
                except Exception as e:
                    print e.message
            paragraph.text = content
    # 读取表格数据
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = "".join([run.text for run in paragraph.runs])
                    properties = get_param(text)
                    if properties:
                        run = paragraph.runs[0]
                        for run in paragraph.runs:
                            run.text = ""
                        for prop in properties:
                            try:
                                _value = eval(prop).decode("utf-8")
                                text = text.replace("{"+prop+"}", _value)
                            except Exception as e:
                                print e.message
                        run.text = text
    print "over"
    document.save("test"+str(time.time())+".docx")

def main1(file_path):
    # 定义所有需要用到的变量
    # 
    #
    document = docx.Document(file_path)
    # 读取每个段落
    for paragraph in document.paragraphs:
        content = paragraph.text
        properties = get_param(content)
        if properties:
            for prop in properties:
                try:
                    _value = eval(prop)
                    if type(_value).__name__ not in ["str", "unicode"]:
                        _value = str(_value)
                    content = content.replace("{"+prop+"}", _value)
                    # content = content.replace("{"+prop+"}", eval(prop))
                except UnicodeDecodeError:
                    content = content.replace("{"+prop+"}", _value.decode("utf-8"))
                except NameError:
                    print("%s is not define, please check it" % prop)
                except Exception as e:
                    print(e.message)
            try:
                paragraph.text = content
            except ValueError:
                paragraph.text = content.decode("utf-8")
            except Exception as e:
                print(e.message)
    # 读取表格数据
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    content = "".join([run.text for run in paragraph.runs])
                    properties = get_param(content)
                    if properties:
                        for run in paragraph.runs:
                            run.text = ""
                        for prop in properties:
                            try:
                                _value = eval(prop)
                                if type(_value).__name__ not in ["str", "unicode"]:
                                    _value = str(_value)
                                content = content.replace("{"+prop+"}", _value)
                            except UnicodeDecodeError:
                                content = content.replace("{"+prop+"}", _value.decode("utf-8"))
                            except NameError:
                                print("%s is not define, please check it" % prop)
                            except Exception as e:
                                print(e.message)
                        try:
                            run.text = content
                        except ValueError:
                            run.text = content.decode("utf-8")
                        except Exception as e:
                            print(e.message)
    print "over"
    document.save("test"+str(time.time())+".docx")



def get_param(string):
    result = re.findall(r"{\w+}", string)
    params = [r.strip("{}") for r in result]
    return params


if __name__ == '__main__':
    try:
        main1(r"D:\code\tk\dainihai_general_product.docx")
    except Exception as e:
        print traceback.format_exc()
    else:
        pass
    finally:
        pass
    

